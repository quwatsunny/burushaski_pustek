import os
import base64
from io import BytesIO
from ebooklib import epub
from docx import Document

def export_epub_py(book_title, chapters):
    book = epub.EpubBook()
    book.set_title(book_title)
    book.set_language('en')
    epub_chapters = []
    for idx, chapter in enumerate(chapters):
        c = epub.EpubHtml(title=chapter['title'], file_name=f'chap_{idx+1}.xhtml', content=chapter['content'])
        book.add_item(c)
        epub_chapters.append(c)
    book.toc = tuple(epub_chapters)
    book.spine = ['nav'] + epub_chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    buf = BytesIO()
    epub.write_epub(buf, book)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')

def export_word_py(book_title, chapters):
    doc = Document()
    doc.add_heading(book_title, 0)
    for chapter in chapters:
        doc.add_heading(chapter['title'], level=1)
        # Naive HTML to text (for demo)
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(chapter['content'], 'html.parser')
        for p in soup.find_all(['p', 'div', 'li']):
            doc.add_paragraph(p.get_text())
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode('utf-8')
